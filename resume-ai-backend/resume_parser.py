# resume_parser.py
import io
import re
from typing import Optional

try:
    import fitz  # PyMuPDF
except Exception as e:
    raise RuntimeError("PyMuPDF is required. Install with: pip install PyMuPDF") from e

from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for page in doc:
        pages.append(page.get_text("text"))
    return "\n".join(pages)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    bio = io.BytesIO(file_bytes)
    doc = Document(bio)
    parts = [p.text for p in doc.paragraphs]
    return "\n".join(parts)


def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text(file_bytes: bytes, filename: Optional[str] = None) -> str:
    """Auto-detect file type from filename and extract text accordingly."""
    fname = (filename or "").lower()
    if fname.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    if fname.endswith(".docx") or fname.endswith(".doc"):
        # Document() can't open legacy .doc; but try anyway
        return extract_text_from_docx(file_bytes)
    if fname.endswith(".txt"):
        return extract_text_from_txt(file_bytes)

    # Fallback: try PDF first, then decode as text
    try:
        return extract_text_from_pdf(file_bytes)
    except Exception:
        return extract_text_from_txt(file_bytes)


# Small helper utilities used by analysis
STOPWORDS = {
    "and", "or", "the", "a", "an", "to", "for", "in", "of", "on", "with", "is", "are",
    "be", "as", "by", "we", "you", "your", "our", "this", "that", "these", "those"
}


def tokenize(text: str):
    # lowercase, remove excessive whitespace
    text = text.lower()
    tokens = re.findall(r"[a-z0-9\+\#\.\-]+", text)
    return tokens


def ngrams(tokens, n):
    return [" ".join(tokens[i:i + n]) for i in range(len(tokens) - n + 1)]
